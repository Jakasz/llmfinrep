"""Word (.docx) document text extraction using python-docx.

Extracts all paragraphs and tables.
"""

import io
import logging

from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a Word document.

    Args:
        file_bytes: Raw .docx file content.

    Returns:
        Text from all paragraphs and tables.
    """
    doc = Document(io.BytesIO(file_bytes))
    parts = []

    # Витягнути текст з усіх параграфів
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Витягнути текст з таблиць
    for table_idx, table in enumerate(doc.tables):
        table_lines = [f"--- Table {table_idx + 1} ---"]
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            table_lines.append("\t".join(cells))
        parts.append("\n".join(table_lines))

    result = "\n".join(parts)
    logger.info(
        "DOCX extraction completed: %d paragraphs, %d tables, %d characters",
        len(doc.paragraphs),
        len(doc.tables),
        len(result),
    )
    return result
